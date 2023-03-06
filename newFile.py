# import openpyxl module
import openpyxl
# from PIL import ImageGrab
import os
import regex as re

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
import docx2pdf as pd
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.font.bold = False
    r.font.italic = False
    r.font.underline = True

    return hyperlink


subdiv = ['Jammu & Kashmir & Ladakh', 'Himachal Pradesh', 'Punjab', 'Uttarakhand', 'Har - Chd & Delhi',
          'West Uttar Pradesh', 'East Uttar Pradesh',
          'West Rajasthan', 'East Rajasthan', 'Saurashtra & Kutch', 'Gujarat Region', 'West Madhya Pradesh',
          'East Madhya Pradesh', 'Konkan & Goa',
          'Madhya Maharashtra', 'Marathawada', 'Vidarbha', 'Telangana', 'NI Karnataka', 'Coastal Karnataka',
          'SI Karnataka', 'Kerala & Mahe',
          'Tamil Nadu - Pudu & Karaikal', 'Rayalaseema', 'Coastal AP & Yanam', 'Chhattisgarh', 'Odisha', 'Jharkhand',
          'Bihar', 'Gangetic West Bengal',
          'SHWB & Sikkim', 'NMMT', 'Assam & Meghalaya', 'Arunachal Pradesh', 'Lakshadweep', 'A & N Island']

intensity = ['Light flash flood threat', 'Light to Moderate flash flood threat', 'Moderate flash flood threat',
             'Moderate to High flash flood threat', 'High flash flood threat', 'Light flash flood risk',
             'Light to Moderate flash flood risk',
             'Moderate flash flood risk', 'Moderate to High flash flood risk', 'High flash flood risk']

svdir = os.getcwd()
path = svdir + "\FFGS Bulletin.xlsm"

wb_obj = openpyxl.load_workbook(path)
sheet_obj = wb_obj.active

bltn_heading = sheet_obj.cell(row=1, column=2).value

dt = sheet_obj.cell(row=3, column=3).value
date_value = sheet_obj.cell(row=3, column=2).value + dt.strftime("%d-%m-%Y")
time_value = sheet_obj.cell(row=3, column=5).value + sheet_obj.cell(row=3, column=6).value
validity = sheet_obj.cell(row=3, column=8).value + sheet_obj.cell(row=3, column=9).value

from_value = sheet_obj.cell(row=5, column=2).value + sheet_obj.cell(row=5, column=3).value
from_hyperlink = sheet_obj.cell(row=5, column=6).value
to_value = sheet_obj.cell(row=7, column=2).value + sheet_obj.cell(row=7, column=3).value
aoc_value_head = sheet_obj.cell(row=9, column=2).value
aoc_value = sheet_obj.cell(row=9, column=4).value

diag_guidance_head = sheet_obj.cell(row=11, column=2).value
issueTime = sheet_obj.cell(row=11, column=6).value
diag_guidance = sheet_obj.cell(row=11, column=4).value + sheet_obj.cell(row=11, column=6).value + \
                sheet_obj.cell(row=12, column=4).value + str(sheet_obj.cell(row=12, column=5).value) + sheet_obj.cell(
    row=12, column=6).value + \
                sheet_obj.cell(row=13, column=4).value + str(sheet_obj.cell(row=13, column=5).value) + sheet_obj.cell(
    row=13, column=6).value + \
                sheet_obj.cell(row=14, column=4).value + sheet_obj.cell(row=14, column=6).value + sheet_obj.cell(row=16,
                                                                                                                 column=4).value + sheet_obj.cell(
    row=16, column=7).value + sheet_obj.cell(row=16, column=8).value + \
                sheet_obj.cell(row=17, column=4).value + sheet_obj.cell(row=19, column=4).value + str(
    sheet_obj.cell(row=19, column=5).value) + \
                sheet_obj.cell(row=19, column=6).value

prog_guidance_head = sheet_obj.cell(row=21, column=2).value
prog_guidance = sheet_obj.cell(row=21, column=4).value + \
                sheet_obj.cell(row=22, column=4).value + sheet_obj.cell(row=22, column=5).value + \
                str(sheet_obj.cell(row=22, column=6).value) + sheet_obj.cell(row=22, column=7).value

# IFFT Data
ifft_date = sheet_obj.cell(row=25, column=7).value
ifft_text = sheet_obj.cell(row=25, column=2).value + sheet_obj.cell(row=25, column=5).value + \
            sheet_obj.cell(row=25, column=6).value + ifft_date.strftime("%d-%m-%Y") + sheet_obj.cell(row=25,
                                                                                                     column=8).value

ifft_rows = sheet_obj.cell(row=27, column=4).value

i1Int = sheet_obj.cell(row=29, column=3).value or ""
i1Val = sheet_obj.cell(row=29, column=4).value or ""

i2Int = sheet_obj.cell(row=30, column=3).value or ""
i2Val = sheet_obj.cell(row=30, column=4).value or ""

i3Int = sheet_obj.cell(row=31, column=3).value or ""
i3Val = sheet_obj.cell(row=31, column=4).value or ""

i4Int = sheet_obj.cell(row=32, column=3).value or ""
i4Val = sheet_obj.cell(row=32, column=4).value or ""
# print(ifft_rows)
ifft1 = "";
ifft2 = "";
ifft3 = "";
ifft4 = "";
i_rw = 0;
ifft = ""
if ifft_rows == "0 (No IFFT)":
    i_rw = 1; ifft1 = sheet_obj.cell(row=28, column=4).value; ifft = "Nil"
elif ifft_rows == 1:
    i_rw = 1; ifft1 = i1Int + "\n\n" + i1Val
elif ifft_rows == 2:
    i_rw = 2; ifft1 = i1Int + "\n\n" + i1Val; ifft2 = i2Int + "\n\n" + i2Val
elif ifft_rows == 3:
    i_rw = 3; ifft1 = i1Int + "\n\n" + i1Val; ifft2 = i2Int + "\n\n" + i2Val; ifft3 = i3Int + "\n\n" + i3Val
elif ifft_rows == 4:
    i_rw = 4; ifft1 = i1Int + "\n\n" + i1Val; ifft2 = i2Int + "\n\n" + i2Val; ifft3 = i3Int + "\n\n" + i3Val; ifft4 = i4Int + "\n\n" + i4Val

# PFFT Data
pfft_date = sheet_obj.cell(row=33, column=7).value
pfft_text = sheet_obj.cell(row=33, column=2).value + sheet_obj.cell(row=33, column=5).value + \
            sheet_obj.cell(row=33, column=6).value + pfft_date.strftime("%d-%m-%Y") + sheet_obj.cell(row=33,
                                                                                                     column=8).value

pfft_rows = sheet_obj.cell(row=35, column=4).value
# print(pfft_rows)
p1Int = sheet_obj.cell(row=37, column=3).value or ""
p1Val = sheet_obj.cell(row=37, column=4).value or ""

p2Int = sheet_obj.cell(row=38, column=3).value or ""
p2Val = sheet_obj.cell(row=38, column=4).value or ""

p3Int = sheet_obj.cell(row=39, column=3).value or ""
p3Val = sheet_obj.cell(row=39, column=4).value or ""

p4Int = sheet_obj.cell(row=40, column=3).value or ""
p4Val = sheet_obj.cell(row=40, column=4).value or ""
# print(ifft_rows)

pfft1 = "";
pfft2 = "";
pfft3 = "";
pfft4 = "";
p_rw = 0;
pfft = ""
if pfft_rows == "0 (No PFFT)":
    p_rw = 1; pfft1 = sheet_obj.cell(row=36, column=4).value; pfft = "Nil"
elif pfft_rows == 1:
    p_rw = 1; pfft1 = p1Int + "\n\n" + p1Val
elif pfft_rows == 2:
    p_rw = 2; pfft1 = p1Int + "\n\n" + p1Val; pfft2 = p2Int + "\n\n" + p2Val
elif pfft_rows == 3:
    p_rw = 3; pfft1 = p1Int + "\n\n" + p1Val; pfft2 = p2Int + "\n\n" + p2Val; pfft3 = p3Int + "\n\n" + p3Val
elif pfft_rows == 4:
    p_rw = 4; pfft1 = p1Int + "\n\n" + p1Val; pfft2 = p2Int + "\n\n" + p2Val; pfft3 = p3Int + "\n\n" + p3Val; pfft4 = p4Int + "\n\n" + p4Val

pfft_surface_txt = sheet_obj.cell(row=43, column=3).value
# FFR Data
ffr_date = sheet_obj.cell(row=46, column=7).value
ffr_text = sheet_obj.cell(row=46, column=2).value + sheet_obj.cell(row=46, column=5).value + \
           sheet_obj.cell(row=46, column=6).value + ffr_date.strftime("%d-%m-%Y") + sheet_obj.cell(row=46,
                                                                                                   column=8).value

ffr_rows = sheet_obj.cell(row=48, column=4).value
# print(ffr_rows)

f1Int = sheet_obj.cell(row=50, column=3).value or ""
f1Val = sheet_obj.cell(row=50, column=4).value or ""

f2Int = sheet_obj.cell(row=51, column=3).value or ""
f2Val = sheet_obj.cell(row=51, column=4).value or ""

f3Int = sheet_obj.cell(row=52, column=3).value or ""
f3Val = sheet_obj.cell(row=52, column=4).value or ""

f4Int = sheet_obj.cell(row=53, column=3).value or ""
f4Val = sheet_obj.cell(row=53, column=4).value or ""

ffr1 = "";
ffr2 = "";
ffr3 = "";
ffr4 = "";
f_rw = 0;
ffr = ""
if ffr_rows == "0 (No FFR)":
    f_rw = 1; ffr1 = sheet_obj.cell(row=49, column=4).value; ffr = "Nil"
elif ffr_rows == 1:
    f_rw = 1; ffr1 = f1Int + f1Val
elif ffr_rows == 2:
    f_rw = 2; ffr1 = f1Int + f1Val; ffr2 = f2Int + f2Val
elif ffr_rows == 3:
    f_rw = 3; ffr1 = f1Int + f1Val; ffr2 = f2Int + f2Val; ffr3 = f3Int + f3Val
elif ffr_rows == 4:
    f_rw = 4; ffr1 = f1Int + f1Val; ffr2 = f2Int + f2Val; ffr3 = f3Int + f3Val; ffr4 = f4Int + f4Val

ffr_surface_txt = sheet_obj.cell(row=56, column=3).value

nxt_bltn_date = sheet_obj.cell(row=59, column=7).value
nxt_bltn_text = sheet_obj.cell(row=59, column=2).value + sheet_obj.cell(row=59, column=5).value + \
                sheet_obj.cell(row=59, column=6).value + nxt_bltn_date.strftime("%d-%m-%Y") + sheet_obj.cell(row=59,
                                                                                                             column=8).value
# Print value of cell object
# using the value attribute
# print(bltn_heading,date_value,time_value,validity,from_value,to_value,aoc_value,diag_guidance,prog_guidance,ifft_text, sep="\n")

document = Document()
section = document.sections[0]

document.sections[0].left_margin = Cm(0.85)
document.sections[0].right_margin = Cm(1.02)
document.sections[0].header_distance = Cm(0)
document.sections[0].footer_distance = Cm(0)
# Add header to Doc
header = section.header
ph = header.paragraphs[0]
run = ph.add_run()
run.add_picture("Images/DND Images/header.jpg")
run.underline = False

# Add footer to Doc
footer = section.footer
pf = footer.paragraphs[0]
run = pf.add_run()
run.add_picture("Images/DND Images/footer.jpg")
run.underline = False

# Add heading to Doc
p = document.add_heading(bltn_heading, 1)
p.alignment = 1  # 0 for left, 1 for center, 2 right, 3 justify ....

p.style = document.styles['Normal']
font = p.style.font
font.name = 'Times New Roman';
font.size = Pt(18);
font.color.rgb = RGBColor(0, 0, 0);
font.bold = True;
font.italic = False;
font.underline = True

# Add table to Doc for Date Time & Validity
table = document.add_table(1, 3)

table.cell(0, 0).text = date_value;
table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
table.cell(0, 1).text = time_value;
table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
table.cell(0, 2).text = validity;
table.cell(0, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT

for row in table.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font;
                font.size = Pt(14);
                font.name = 'Times New Roman';
                font.color.rgb = RGBColor(0, 0, 0)
                font.bold = True;
                font.italic = False;
                font.underline = False

                # Add new paragraph to Doc
p1 = document.add_paragraph(from_value)
to_mail = "sasiaffg.imd@gmail.com"

mail_to_link = f"mailto:{to_mail}"

# Adding the mail to link as any other hyperlink
add_hyperlink(p1, 'sasiaffg.imd@gmail.com', mail_to_link)
run_p1 = p1.add_run(")")
run_p1.font.size = Pt(14)
run_p1.font.bold = False

p1.alignment = 3
font = p1.runs[0].font
font.name = 'Times New Roman';
font.size = Pt(14);
font.color.rgb = RGBColor(0, 0, 0);
font.bold = False;
font.italic = False;
font.underline = False

# Add new paragraph to Doc
p2 = document.add_paragraph(to_value)
p2.alignment = 3
font = p2.runs[0].font
font.name = 'Times New Roman';
font.size = Pt(14);
font.color.rgb = RGBColor(0, 0, 0);
font.bold = True;
font.italic = False;
font.underline = False

# Add new paragraph and add text to the paragraph to Doc
p3 = document.add_paragraph(aoc_value_head)
p3.alignment = 3
font = p3.runs[0].font
font.name = 'Times New Roman';
font.size = Pt(14);
font.color.rgb = RGBColor(0, 0, 0);
font.bold = True;
font.italic = False;
font.underline = True
# Adding text to above paragraph
bold_para = p3.add_run(aoc_value)
font = bold_para.font
font.name = 'Times New Roman';
font.size = Pt(14);
font.color.rgb = RGBColor(0, 0, 0);
font.bold = False;
font.italic = False;
font.underline = False

# Add new paragraph and add text to the paragraph to Doc
p4 = document.add_paragraph(diag_guidance_head)
p4.alignment = 3
font = p4.runs[0].font
font.name = 'Times New Roman';
font.size = Pt(14);
font.color.rgb = RGBColor(0, 0, 0);
font.bold = True;
font.italic = False;
font.underline = True
# Adding text to above paragraph
bold_para1 = p4.add_run(diag_guidance)
font = bold_para1.font
font.name = 'Times New Roman';
font.size = Pt(14);
font.color.rgb = RGBColor(0, 0, 0);
font.bold = False;
font.italic = False;
font.underline = False

# Add new paragraph and add text to the paragraph to Doc
p5 = document.add_paragraph(prog_guidance_head)
p5.alignment = 3
font = p5.runs[0].font
font.name = 'Times New Roman';
font.size = Pt(14);
font.color.rgb = RGBColor(0, 0, 0);
font.bold = True;
font.italic = False;
font.underline = True
# Adding text to above paragraph
bold_para2 = p5.add_run(prog_guidance)
font = bold_para2.font
font.name = 'Times New Roman';
font.size = Pt(14);
font.color.rgb = RGBColor(0, 0, 0);
font.bold = False;
font.italic = False;
font.underline = False

# Add table to Doc for IFFT
table1 = document.add_table(i_rw, 2)
# print(i_rw)
table1.style = 'Table Grid'
for i in range(i_rw):
    table1.cell(i, 0).text = ifft_text
    table1.cell(i, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
    table1.cell(i, 0).paragraphs[0].runs[0].font.size = Pt(14)
    nm_ifft = "";
    img_ifft = ""
    if ifft != "Nil":
        if i == 0:
            nm_ifft = ifft1; img_ifft = "Images\ifft1.png"
        elif i == 1:
            nm_ifft = ifft2; img_ifft = "Images\ifft2.png"
        elif i == 2:
            nm_ifft = ifft3; img_ifft = "Images\ifft3.png"
        elif i == 3:
            nm_ifft = ifft4; img_ifft = "Images\ifft4.png"

    if i_rw == 1 and ifft == "Nil":
        p_i = table1.cell(i, 0).add_paragraph("\n" + ifft1)
        pi_font = p_i.runs[0].font
        pi_font.name = 'Times New Roman'
        pi_font.size = Pt(14)
        pi_font.bold = False
        pi_font.italic = False
        pi_font.underline = False

        run1 = table1.cell(i, 1).paragraphs[0].add_run()
        run1.add_picture("Images/DND Images/No IFFT.png", width=Cm(9.66), height=Cm(7.81))
        run1.underline = False
    else:
        # print("\n" + nm_ifft)
        int_ifft = re.split(rf"({'|'.join(intensity)})", nm_ifft)
        subdivi = re.split(rf"({'|'.join(subdiv)})", int_ifft[2])
        # print(len(subdivi),subdivi[0], subdivi[1])

        p_i = table1.cell(i, 0).add_paragraph("\n" + int_ifft[1])
        p_i.alignment = 3
        pi_font = p_i.runs[0].font
        pi_font.name = 'Times New Roman'
        pi_font.size = Pt(14)
        pi_font.bold = True
        pi_font.italic = False
        pi_font.underline = False

        for si in range(len(subdivi)):
            # print(subdivi[si])
            if si % 2 == 0:
                run_pi = p_i.add_run(subdivi[si])
            else:
                run_pi = p_i.add_run("\n" + subdivi[si])
            run_pi_font = p_i.runs[si + 1].font
            run_pi_font.name = 'Times New Roman'
            run_pi_font.size = Pt(14)
            run_pi_font.italic = False
            run_pi_font.underline = False
            if si % 2 == 0:
                run_pi_font.bold = False
            else:
                run_pi_font.bold = True

        try:
            run2 = table1.cell(i, 1).paragraphs[0].add_run()
            run2.add_picture(img_ifft, width=Cm(9.66), height=Cm(7.81))
        except:
            run2 = table1.cell(i, 1).paragraphs[0].add_run("Image Not Found")
        run2.underline = False

        # Add table to Doc for PFFT
table2 = document.add_table(p_rw, 2)
# print(p_rw)
table2.style = 'Table Grid'
for i in range(p_rw):
    table2.cell(i, 0).text = pfft_text
    table2.cell(i, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
    table2.cell(i, 0).paragraphs[0].runs[0].font.size = Pt(14)
    nm_pfft = "";
    img_pfft = ""
    if pfft != "Nil":
        if i == 0:
            nm_pfft = pfft1; img_pfft = "Images\pfft1.png"
        elif i == 1:
            nm_pfft = pfft2; img_pfft = "Images\pfft2.png"
        elif i == 2:
            nm_pfft = pfft3; img_pfft = "Images\pfft3.png"
        elif i == 3:
            nm_pfft = pfft4; img_pfft = "Images\pfft4.png"

    if p_rw == 1 and pfft == "Nil":
        p_i = table2.cell(i, 0).add_paragraph("\n" + pfft1)
        pi_font = p_i.runs[0].font
        pi_font.name = 'Times New Roman'
        pi_font.size = Pt(14)
        pi_font.bold = False
        pi_font.italic = False
        pi_font.underline = False

        run2 = table2.cell(i, 1).paragraphs[0].add_run()
        run2.add_picture("Images/DND Images/No PFFT.png", width=Cm(9.66), height=Cm(7.81))
        run2.underline = False
    else:
        # print("\n" + nm_pfft + "\n\n" + pfft_surface_txt)
        int_pfft = re.split(rf"({'|'.join(intensity)})", nm_pfft)
        subdivi = re.split(rf"({'|'.join(subdiv)})", int_pfft[2])
        # print(len(subdivi), subdivi[0], subdivi[1])

        p_i = table2.cell(i, 0).add_paragraph("\n" + int_pfft[1])
        p_i.alignment = 3
        pi_font = p_i.runs[0].font
        pi_font.name = 'Times New Roman'
        pi_font.size = Pt(14)
        pi_font.bold = True
        pi_font.italic = False
        pi_font.underline = False

        for si in range(len(subdivi)):
            # print(subdivi[si])
            if si % 2 == 0:
                run_pi = p_i.add_run(subdivi[si])
            else:
                run_pi = p_i.add_run("\n" + subdivi[si])
            run_pi_font = p_i.runs[si + 1].font
            run_pi_font.name = 'Times New Roman'
            run_pi_font.size = Pt(14)
            run_pi_font.italic = False
            run_pi_font.underline = False
            if si % 2 == 0:
                run_pi_font.bold = False
            else:
                run_pi_font.bold = True

        run_surface = p_i.add_run("\n\n" + pfft_surface_txt)
        run_sur_font = run_surface.font
        run_sur_font.name = 'Times New Roman'
        run_sur_font.size = Pt(14)
        run_sur_font.bold = False
        run_sur_font.italic = False
        run_sur_font.underline = False

        try:
            run2 = table2.cell(i, 1).paragraphs[0].add_run()
            run2.add_picture(img_pfft, width=Cm(9.66), height=Cm(7.81))
        except:
            run2 = table2.cell(i, 1).paragraphs[0].add_run("Image Not Found")
        run2.underline = False

        # Add table to Doc for FFR

table3 = document.add_table(f_rw, 2)
# print(f_rw)
table3.style = 'Table Grid'
for i in range(f_rw):
    table3.cell(i, 0).text = ffr_text
    table3.cell(i, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
    table3.cell(i, 0).paragraphs[0].runs[0].font.size = Pt(14)
    nm_ffr = "";
    img_ffr = ""
    if ffr != "Nil":
        if i == 0:
            nm_ffr = ffr1; img_ffr = r"Images\ffr1.png"
        elif i == 1:
            nm_ffr = ffr2; img_ffr = r"Images\ffr2.png"
        elif i == 2:
            nm_ffr = ffr3; img_ffr = r"Images\ffr3.png"
        elif i == 3:
            nm_ffr = ffr4; img_ffr = r"Images\ffr4.png"

    if f_rw == 1 and ffr == "Nil":
        p_i = table3.cell(i, 0).add_paragraph("\n" + ffr1)
        pi_font = p_i.runs[0].font
        pi_font.name = 'Times New Roman'
        pi_font.size = Pt(14)
        pi_font.bold = False
        pi_font.italic = False
        pi_font.underline = False

        run3 = table3.cell(i, 1).paragraphs[0].add_run()
        run3.add_picture("Images/DND Images/No FFR.png", width=Cm(9.66), height=Cm(7.81))
        run3.underline = False
    else:
        # print("\n" + nm_ffr + "\n" + "\n" + ffr_surface_txt)
        # print(nm_ffr)
        int_ffr = re.split(rf"({'|'.join(intensity)})", nm_ffr)
        # print(len(int_ffr))
        subdivi = re.split(rf"({'|'.join(subdiv)})", int_ffr[2])
        # print(len(subdivi), subdivi[0], subdivi[1])

        p_i = table3.cell(i, 0).add_paragraph("\n" + int_ffr[1])
        p_i.alignment = 3
        pi_font = p_i.runs[0].font
        pi_font.name = 'Times New Roman'
        pi_font.size = Pt(14)
        pi_font.bold = True
        pi_font.italic = False
        pi_font.underline = False

        for si in range(len(subdivi)):

            run_pi = p_i.add_run(subdivi[si])
            run_pi_font = p_i.runs[si + 1].font
            run_pi_font.name = 'Times New Roman'
            run_pi_font.size = Pt(14)
            run_pi_font.italic = False
            run_pi_font.underline = False
            if subdivi[si] in subdiv:
                # print(subdivi[si])
                run_pi_font.bold = True
            else:
                print(subdivi[si])
                run_pi_font.bold = False

        run_surface = p_i.add_run("\n\n" + ffr_surface_txt)
        run_sur_font = run_surface.font
        run_sur_font.name = 'Times New Roman'
        run_sur_font.size = Pt(14)
        run_sur_font.bold = False
        run_sur_font.italic = False
        run_sur_font.underline = False

        try:
            run2 = table3.cell(i, 1).paragraphs[0].add_run()
            run2.add_picture(img_ffr, width=Cm(9.66), height=Cm(7.81))
        except:
            run2 = table3.cell(i, 1).paragraphs[0].add_run("Image Not Found")
        run2.underline = False

# Add new paragraph to Doc
p_next = document.add_paragraph("\n" + nxt_bltn_text)
p_next.alignment = 0
font = p_next.runs[0].font
font.name = 'Times New Roman';
font.size = Pt(14);
font.color.rgb = RGBColor(0, 0, 0);
font.bold = True;
font.italic = False;
font.underline = False

term = document.add_paragraph()
ph1 = term.add_run()
ph1.add_picture("Images/DND Images/terminology.jpg", width=Cm(19.82), height=Cm(20.47))
ph1.underline = False

bltnNameDocx = "IMD National Flash Flood Guidance Bulletin dated " + dt.strftime(
    "%d.%m.%Y") + " (" + issueTime + ").docx"
bltnNamePdf = "IMD National Flash Flood Guidance Bulletin dated " + dt.strftime("%d.%m.%Y") + " (" + issueTime + ").pdf"

document.save(bltnNameDocx)
print(bltnNameDocx, "\n", bltnNamePdf)
print("FFGS Bulletin has been created successfully.")
pd.convert(bltnNameDocx, bltnNamePdf)
pd.convert(bltnNameDocx, "national.pdf")

print("PDFs also created.")

from email.message import Message
from email.mime.multipart import MIMEMultipart, MIMEBase
from email.mime.text import MIMEText
from email.encoders import encode_base64
import imaplib
import time

with imaplib.IMAP4_SSL(host="imap.gmail.com", port=imaplib.IMAP4_SSL_PORT) as imap_ssl:
    print("Logging into mailbox...")
    resp_code, response = imap_ssl.login('dgmfmu@gmail.com', 'ccpr kkro bkni bxkm')

    email_message = MIMEMultipart()
    email_message.add_header('To', 'yashika.garg23@gmail.com')
    email_message.add_header('From', 'dgmfmu@gmail.com')
    email_message.add_header('Subject', bltnNameDocx[:-5])
    email_message.add_header('X-Priority', '1')  # Urgent/High priority

    # Create text and HTML bodies for email
    html_part = MIMEText("<html><body><h4>Respected Madam/Sir,</h4></body></html>", "html")
    html_part1 = MIMEText("<html><body><h3>\nKindly find attached " + bltnNameDocx[:-5] + "</h3></body></html>", "html")
    html_part2 = MIMEText("<html><body><h3>" + aoc_value_head + aoc_value + "</h3></body></html>", "html")
    # Create file attac
    # hment
    attachment = MIMEBase("application", "octet-stream")
    attachment.set_payload(open(bltnNamePdf, "rb").read())  # Raw attachment data
    encode_base64(attachment)
    attachment.add_header('Content-Disposition', 'attachment', filename=bltnNamePdf)

    # Attach all the parts to the Multipart MIME email
    email_message.attach(html_part)
    email_message.attach(html_part1)
    email_message.attach(html_part2)
    email_message.attach(attachment)

    utf8_message = str(email_message).encode("utf-8")

    # Send message
    imap_ssl.append("[Gmail]/Drafts", '', imaplib.Time2Internaldate(time.time()), utf8_message)

input('Press Enter to exit')